from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from typing import List, Dict, Any
from datetime import datetime
import logging

logger = logging.getLogger(__name__)

class DocxReportGenerator:
    @staticmethod
    async def generate_report(
        test: Dict[str, Any],
        session: Dict[str, Any],
        answers: List[Dict[str, Any]],
        psychologist: Dict[str, Any]
    ) -> BytesIO:
        """Генерирует отчет в формате DOCX."""
        doc = Document()
        
        # Заголовок
        title = doc.add_heading(f'Отчет по тесту: {test["title"]}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информация о клиенте
        doc.add_heading('Информация о клиенте', level=1)
        doc.add_paragraph(f'Имя: {session["client_name"]}')
        if session.get("client_email"):
            doc.add_paragraph(f'Email: {session["client_email"]}')
        if session.get("client_phone"):
            doc.add_paragraph(f'Телефон: {session["client_phone"]}')
        
        doc.add_paragraph()  # Отступ
        
        # Информация о психологе
        doc.add_heading('Информация о психологе', level=1)
        doc.add_paragraph(f'Психолог: {psychologist.get("full_name", "Не указано")}')
        doc.add_paragraph(f'Email: {psychologist.get("email", "Не указан")}')
        
        doc.add_paragraph()  # Отступ
        
        # Информация о тесте
        doc.add_heading('Информация о тестировании', level=1)
        doc.add_paragraph(f'Название: {test.get("title", "—")}')
        if test.get("description"):
            doc.add_paragraph(f'Описание: {test["description"]}')
        doc.add_paragraph(f'Дата прохождения: {session.get("completed_at", "Не указана")}')
        doc.add_paragraph(f'Статус: {session.get("status", "Завершен")}')
        
        doc.add_paragraph()  # Отступ
        
        # Результаты
        doc.add_heading('Результаты тестирования', level=1)
        
        for answer in answers:
            # Вопрос
            p = doc.add_paragraph()
            p.add_run(f'Вопрос {answer["order_index"] + 1}: ').bold = True
            p.add_run(answer["question_text"])
            
            # Ответ
            answer_text = answer["answer_value"]
            if isinstance(answer_text, list):
                answer_text = ", ".join(answer_text)
            elif isinstance(answer_text, dict):
                answer_text = str(answer_text)
            
            doc.add_paragraph(f'Ответ: {answer_text}', style='List Bullet')
            doc.add_paragraph()  # Пустая строка после ответа
        
        # Дата генерации
        doc.add_paragraph(f'Отчет сгенерирован: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        
        # Сохраняем в BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info(f"Report generated for session {session.get('id')}")
        return buffer